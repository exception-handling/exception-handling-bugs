"""
Generates eh_mechanisms_val/INSTRUCTIONS_FOR_AUTHORS.docx
"""

import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = "eh_mechanisms_val/INSTRUCTIONS_FOR_AUTHORS.docx"

doc = Document()

section = doc.sections[0]
section.top_margin    = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin   = Inches(1.2)
section.right_margin  = Inches(1.2)


def h1(text):
    p = doc.add_heading(text, level=1)
    p.runs[0].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    return p

def h2(text):
    p = doc.add_heading(text, level=2)
    p.runs[0].font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
    return p

def h3(text):
    p = doc.add_heading(text, level=3)
    p.runs[0].font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
    return p

def _fill_runs(p, text, bold_parts):
    if not bold_parts:
        p.add_run(text)
        return
    remaining = text
    for bp in bold_parts:
        idx = remaining.find(bp)
        if idx == -1:
            continue
        if idx > 0:
            p.add_run(remaining[:idx])
        p.add_run(bp).bold = True
        remaining = remaining[idx + len(bp):]
    if remaining:
        p.add_run(remaining)

def body(text, bold_parts=None):
    p = doc.add_paragraph()
    p.style = doc.styles["Normal"]
    _fill_runs(p, text, bold_parts)
    return p

def note(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.italic = True
    run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
    return p

def bullet(text, bold_parts=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.3)
    _fill_runs(p, text, bold_parts)
    return p

def numbered(text, bold_parts=None):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.left_indent = Inches(0.3)
    _fill_runs(p, text, bold_parts)
    return p

def code_inline(paragraph, text):
    run = paragraph.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(9.5)
    return run

def spacer():
    doc.add_paragraph()

def shade_cell(cell, fill_hex):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill_hex)
    shd.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shd)

def add_table(headers, rows, col_widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    hrow = t.rows[0]
    for i, h in enumerate(headers):
        cell = hrow.cells[i]
        cell.text = h
        run = cell.paragraphs[0].runs[0]
        run.bold = True
        run.font.size = Pt(9.5)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), "2E74B5")
        shd.set(qn("w:color"), "FFFFFF")
        shd.set(qn("w:val"), "clear")
        cell._tc.get_or_add_tcPr().append(shd)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for ri, row in enumerate(rows):
        trow = t.rows[ri + 1]
        for ci, val in enumerate(row):
            cell = trow.cells[ci]
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            parts = re.split(r"(`[^`]+`)", val)
            for part in parts:
                if part.startswith("`") and part.endswith("`"):
                    r = p.add_run(part[1:-1])
                    r.font.name = "Courier New"
                    r.font.size = Pt(9)
                else:
                    r = p.add_run(part)
                    r.font.size = Pt(9.5)
        if ri % 2 == 1:
            for ci in range(len(row)):
                shade_cell(trow.cells[ci], "EBF3FB")
    if col_widths:
        for row in t.rows:
            for ci, w in enumerate(col_widths):
                row.cells[ci].width = Inches(w)
    return t


# ===========================================================================
# Document
# ===========================================================================

h1("Validation Instructions - EH Mechanisms Extraction")

body(
    "You have two tasks to complete, each with its own CSV file. "
    "Work independently - do not share your answers with the other author "
    "until both of you are done with both parts.",
    bold_parts=["two tasks", "Work independently"],
)

# --- Part A -----------------------------------------------------------------
doc.add_page_break()
h2("Part A - Positive Sample (mechanism was detected)")

p = doc.add_paragraph()
p.add_run("File: ").bold = True
code_inline(p, "eh_mechanisms_sample_authorX.csv")
p.add_run("  (384 rows)")

spacer()
body(
    "Each row is a mechanism the miner claimed to detect. "
    "Your job: confirm the extracted snippet matches the claimed type.",
    bold_parts=["confirm the extracted snippet matches the claimed type."],
)

spacer()
h3("Columns to read")

add_table(
    headers=["Column", "What it contains"],
    rows=[
        ["`mechanism_type`",
         "The type the miner claimed to have detected"],
        ["`mechanism_snippet`",
         "The actual code snippet extracted by the miner"],
        ["`func_body`",
         "The full function body - use for context if the snippet is unclear"],
    ],
    col_widths=[1.8, 4.2],
)

spacer()
h3("Decision criteria - is_correct")

add_table(
    headers=["Mechanism type", "Mark Y if...", "Mark N if..."],
    rows=[
        ["`try-except`",
         "The snippet contains a `try` block with at least one `except` clause",
         "No `except` clause is present"],
        ["`raise`",
         "The snippet is a `raise` statement (with or without an argument)",
         "It is not a raise statement"],
        ["`try-finally`",
         "The snippet contains a `try` block with a `finally` clause",
         "No `finally` clause is present"],
        ["`try-else`",
         "The snippet contains a `try` block with both an `except` and an `else` clause",
         "Either clause is missing"],
    ],
    col_widths=[1.4, 2.8, 2.4],
)

spacer()
note(
    "You are not judging code quality, style, or best practices. "
    "Only check whether the snippet matches the claimed mechanism_type."
)

spacer()
h3("How to fill in Part A")

bullet("Write Y or N in is_correct.", bold_parts=["Y", "N", "is_correct."])
bullet("notes is optional for Y rows.", bold_parts=["notes", "Y"])

p = doc.add_paragraph(style="List Bullet")
p.paragraph_format.left_indent = Inches(0.3)
p.add_run("For ")
p.add_run("N").bold = True
p.add_run(" rows, ")
p.add_run("always").bold = True
p.add_run(' add a brief explanation (e.g. "snippet is a try-finally, not try-except").')


# --- Part B -----------------------------------------------------------------
doc.add_page_break()
h2("Part B - Negative Sample (miner detected nothing)")

p = doc.add_paragraph()
p.add_run("File: ").bold = True
code_inline(p, "eh_negative_sample_authorX.csv")
p.add_run("  (384 rows)")

spacer()
body(
    "Each row is a function where the miner found no exception handling "
    "mechanism. There is no snippet - the miner predicted nothing was there. "
    "Your job: check whether the function actually contains any EH mechanism "
    "that the miner missed.",
    bold_parts=[
        "no exception handling mechanism",
        "check whether the function actually contains any EH mechanism "
        "that the miner missed.",
    ],
)

spacer()
h3("Column to read")

add_table(
    headers=["Column", "What it contains"],
    rows=[
        ["`func_body`", "The full function body to inspect"],
    ],
    col_widths=[1.8, 4.2],
)

spacer()
h3("Decision criteria - has_mechanism")

body("Read func_body carefully and look for any of the four constructs:",
     bold_parts=["func_body"])
spacer()

add_table(
    headers=["Construct", "What to look for"],
    rows=[
        ["`try-except`",  "A `try:` block with at least one `except` clause"],
        ["`raise`",       "A `raise` statement anywhere in the function"],
        ["`try-finally`", "A `try:` block with a `finally:` clause"],
        ["`try-else`",    "A `try:` block with both an `except` and an `else:` clause"],
    ],
    col_widths=[1.5, 4.5],
)

spacer()

p = doc.add_paragraph()
p.add_run("has_mechanism = N").bold = True
p.add_run(" - no EH construct found -> miner was correct (True Negative)")

p = doc.add_paragraph()
p.add_run("has_mechanism = Y").bold = True
p.add_run(" - you spotted a construct the miner missed -> False Negative")

spacer()
h3("How to fill in Part B")

bullet("Write Y or N in has_mechanism.", bold_parts=["Y", "N", "has_mechanism."])

p = doc.add_paragraph(style="List Bullet")
p.paragraph_format.left_indent = Inches(0.3)
p.add_run("If ")
p.add_run("Y").bold = True
p.add_run(": fill in ")
code_inline(p, "mechanism_found")
p.add_run(" with the type(s) you spotted (e.g. ")
code_inline(p, "raise")
p.add_run(", or ")
code_inline(p, "try-except, raise")
p.add_run(" if you found more than one).")

bullet("notes is optional but helpful when marking Y.",
       bold_parts=["notes", "Y"])


# --- After both authors -----------------------------------------------------
doc.add_page_break()
h2("After both authors are done")

numbered("Share your completed files with each other.")
numbered(
    "Compute inter-rater agreement (Cohen's Kappa) separately for Part A "
    "and Part B - use experiment_val/agreement.py as a reference.",
    bold_parts=["Cohen's Kappa"],
)
numbered(
    "For every row where your answers differ, discuss until you reach consensus.",
    bold_parts=["discuss until you reach consensus."],
)

p = doc.add_paragraph(style="List Number")
p.paragraph_format.left_indent = Inches(0.3)
p.add_run("Record the consensus in both author files, then run: ")
code_inline(p, "python3 eh_mechanisms_val/compute_precision.py")
p.add_run(" to get precision, recall, F1, and accuracy.")


# --- FAQ --------------------------------------------------------------------
doc.add_page_break()
h2("Frequently Asked Questions")

faq = [
    (
        "The snippet is cut off - can I still judge it? (Part A)",
        'Use func_body to find the full construct. If even the function body '
        'is unclear, mark N and note "snippet truncated, cannot confirm".',
    ),
    (
        "The try block has both except and finally - which type is it? (Part A)",
        "A single try block can be two mechanisms. Trust only the "
        "mechanism_type column. If the claimed type is try-except and the "
        "snippet has an except clause (even if it also has a finally), mark Y.",
    ),
    (
        "The raise is inside an except block - does it count? (Parts A & B)",
        "Yes. A raise is a valid mechanism regardless of where it appears "
        "in the function.",
    ),
    (
        "The function is a test or stub - does that matter? (Parts A & B)",
        "No. Validation is about extraction correctness, not code quality. "
        "If the construct is present, mark Y.",
    ),
    (
        "The function body is very short or empty - Part B",
        "If the function body contains only pass, a docstring, or a single "
        "expression with no EH construct, mark has_mechanism = N.",
    ),
]

for question, answer in faq:
    p = doc.add_paragraph()
    p.add_run(question).bold = True
    doc.add_paragraph(answer)
    spacer()


doc.save(OUT)
print("Saved: " + OUT)
